import os
from typing import List, Set

from telethon import TelegramClient
from telethon.tl.types import Message

from link_utils import URL_REGEX, BARE_URL_REGEX, DOMAIN_URL_REGEX, _normalize_url


# ======================
# Settings
# ======================

# ✅ مهم: لا تستخدم /tmp على Render
LOCAL_TMP_DIR = "data/tmp_files"

# ✅ تجاهل الملفات الكبيرة
MAX_FILE_SIZE_MB = 15  # غيرها مثل ما تبغى
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


# ======================
# Public API
# ======================

async def extract_links_from_file(
    client: TelegramClient,
    message: Message
) -> List[str]:
    """
    استخراج الروابط من الملفات بدون ما يضغط /tmp على Render
    - PDF
    - DOCX
    """
    if not message.file:
        return []

    # ✅ Skip large files (مهم جداً)
    try:
        if getattr(message.file, "size", 0) and message.file.size > MAX_FILE_SIZE_BYTES:
            return []
    except Exception:
        pass

    links: Set[str] = set()

    filename = message.file.name or "file"
    mime = (message.file.mime_type or "").lower()

    # ✅ إذا بدون امتداد نحدد حسب mime
    if "." not in filename.lower():
        if mime == "application/pdf":
            filename += ".pdf"
        elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            filename += ".docx"

    # ✅ مجلد تحميل محلي بدل /tmp
    os.makedirs(LOCAL_TMP_DIR, exist_ok=True)
    path = os.path.join(LOCAL_TMP_DIR, filename)

    # لو الاسم يتكرر نغيره
    base, ext = os.path.splitext(path)
    i = 1
    while os.path.exists(path):
        path = f"{base}_{i}{ext}"
        i += 1

    try:
        # تحميل الملف
        await client.download_media(message, path)

        # PDF
        if path.lower().endswith(".pdf") or mime == "application/pdf":
            links.update(_extract_from_pdf(path))

        # DOCX
        elif (
            path.lower().endswith(".docx")
            or mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            links.update(_extract_from_docx(path))

    finally:
        # ✅ حذف الملف مباشرة بعد الاستخراج (أساسي)
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception:
            pass

    # Keep raw links as-is (normalize here only trims)
    links = {_normalize_url(u) for u in links if u}

    return list(links)


# ======================
# Helpers
# ======================

def _extract_urls_from_text(text: str) -> Set[str]:
    found: Set[str] = set()
    if not text:
        return found

    for u in URL_REGEX.findall(text):
        found.add(u)

    for u in BARE_URL_REGEX.findall(text):
        found.add(u)

    for u in DOMAIN_URL_REGEX.findall(text):
        if len(u) >= 6:
            found.add(u)

    return found


# ======================
# PDF
# ======================

def _extract_from_pdf(path: str) -> List[str]:
    links: Set[str] = set()

    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(path)

        for page in reader.pages:
            text = page.extract_text() or ""
            links.update(_extract_urls_from_text(text))

            # annotations hyperlinks
            try:
                annots = page.get("/Annots", [])
                if annots:
                    for a in annots:
                        obj = a.get_object()
                        action = obj.get("/A", None)
                        if action:
                            uri = action.get("/URI", None)
                            if uri:
                                links.add(str(uri))
            except Exception:
                pass

    except Exception:
        pass

    return list(links)


# ======================
# DOCX
# ======================

def _extract_from_docx(path: str) -> List[str]:
    links: Set[str] = set()

    try:
        from docx import Document

        doc = Document(path)

        for para in doc.paragraphs:
            links.update(_extract_urls_from_text(para.text))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    links.update(_extract_urls_from_text(cell.text))

        # Hyperlinks relationships
        try:
            rels = doc.part.rels
            for rel in rels.values():
                if rel.reltype and "hyperlink" in rel.reltype:
                    target = getattr(rel, "target_ref", None)
                    if target:
                        links.add(str(target))
        except Exception:
            pass

    except Exception:
        pass

    return list(links)
